from main import request
from main import requests
from main import render_template
from flask import Blueprint
import pymongo
from main import datetime
from flask import Response
import os
from docx import Document
from docx.shared import Pt, RGBColor
import simplejson as json
import subprocess
import re

blueprint = Blueprint("tcs", __name__, url_prefix="/tcs")


@blueprint.route('/patch', methods=['GET', 'POST'])
def patch():  # Patch: issue에서 버튼 눌러 이동되는 페이지, css 업로드 페이지
    issue_key = request.args.get('issuekey')
    if issue_key is not None:
        conn = pymongo.MongoClient("219.240.43.93", 27017)
        db = conn.tcs

        col = db.user_data
        data2 = list(col.find({}, {"_id": 0, "name": -1, "employee_No": -1, "team": -1}))
        user_list = sorted(data2, reverse=False, key=lambda x: (x['name']))

        col = db.id_pw
        pw_data = col.find({})

        r_tcs = requests.get('https://tcs.telechips.com:8443/rest/api/2/issue/' + issue_key, auth=(pw_data[1]['id'], pw_data[1]['pw']))
        issue_data = json.loads(r_tcs.text)

        project_name = issue_data['fields']['project']['name'].split(']')[-1]

        if issue_data['fields']['customfield_10686'] is not None:  # customfield_10686 = SDK Version
            sdk_version = issue_data['fields']['customfield_10686']
        else:
            sdk_version = 'SDK version 확인 필요'

        patch_type = issue_data['fields']['customfield_11726']['value']  # customfield_11726 = patch type

        if issue_data['fields']['customfield_10685'] is not None:
            patch_version = issue_data['fields']['customfield_10685']  # customfield_10685 = patch version
        else:
            patch_version = 'Patch version 확인 필요'

        summary = issue_data['fields']['summary'].split(']')[-1]

        im_qa_key = issue_data['fields']['customfield_12001']
        css_project_id = ''

        if patch_type is not None:
            if patch_type == 'TIMS (고객사 이슈)':
                if im_qa_key is not None:
                    headers = {"accept": "application/json;charset=UTF-8", "x-token": "123qwe123"}
                    request_tims = requests.get('https://css.telechips.com/api/tims/list?pjtKey=' + im_qa_key.split('-')[0], headers=headers)
                    tims_issue = json.loads(request_tims.text)
                    if "list" not in tims_issue.keys() or len(tims_issue['list']) != 1:
                        css_project_key = 'Error TIMS Issue key 확인 필요'
                    else:
                        css_project_key = tims_issue['list'][0]['pjtNm']
                        css_project_id = tims_issue['list'][0]['pjtId']
                else:
                    css_project_key = 'Patch Issue key 확인 필요'

            elif patch_type == 'TCS (QA 검증 이슈)':
                if im_qa_key is not None:
                    headers = {"accept": "application/json;charset=UTF-8", "x-token": "123qwe123"}
                    request_tcs = requests.get('https://css.telechips.com/api/tcs/list?pjtKey=' + im_qa_key.split('-')[0], headers=headers)
                    tcs_issue = json.loads(request_tcs.text)
                    if "list" not in tcs_issue.keys() or len(tcs_issue['list']) != 1:
                        css_project_key = 'Patch Issue key 확인 필요'

                    else:
                        css_project_key = tcs_issue['list'][0]['pjtNm']
                        css_project_id = tcs_issue['list'][0]['pjtId']
                else:
                    css_project_key = 'Patch Issue key 확인 필요'

            elif patch_type == 'RND (신규 기능 추가/개선)':
                im_qa_key = ''
                css_project_key = ''
                css_project_id = ''

            else:
                css_project_key = ''
                css_project_id = ''

        else:
            patch_type = 'Patch Type 없음 확인 필요'
            css_project_key = 'Patch Issue key 확인 필요'
            css_project_id = ''

        col = db.css_sdk_list
        data = list(col.find({}, {"_id": 0}))
        data_list = sorted(data, reverse=False, key=lambda x: (x['sdk_name']))

        sdk_list_dict = {}

        for i in range(0, len(data_list)):
            sdk_list_dict.setdefault(data_list[i]['sdk_name'], [data_list[i]['sdk_idx'], data_list[i]['parent_idx'], data_list[i]['patch_list']])

        sdk_name_list = list(sdk_list_dict.keys())

        headers = {"accept": "application/json;charset=UTF-8", "x-token": "123qwe123"}

        r_css_tims = requests.get('https://css.telechips.com/api/tims/list', headers=headers)
        css_tims = json.loads(r_css_tims.text)

        css_tims['list']

        conn.close()

        return render_template(
                "git_test.html",
                user=user_list,
                im_qa_key=im_qa_key,
                project_name=project_name,
                project_id=css_project_id,
                tims_project_list=css_tims['list'],
                sdk_version=sdk_version,
                patch_type=patch_type,
                issue_key=issue_key,
                css_project_key=css_project_key,
                patch_version=patch_version,
                summary=summary,
                sdk_list=sdk_list_dict,
                sdk_name=sdk_name_list
            )

    else:
        return '잘못된 접근 입니다!!'


@blueprint.route('/css', methods=['POST'])
def css():  # Patch: CSS에 실제로 patch를 업로드하는 함수
    data = request.form

    patch_description = re.sub('[/$&*?]', '', data['patch_des'])
    sdk_name = ''
    for i in range(2, len(data['sdkIdx'].split('_'))):
        sdk_name += data['sdkIdx'].split('_')[i] + '_'

    f1 = request.files['patchFile']
    f1_ext = f1.filename.split('.')[1]
    f1.save('/home/B180093/anaconda3/envs/telechips/main/static/download/' + sdk_name[:-1] + '_' + data['patch_version'] + '_' + patch_description + '.' + f1_ext)
    files1 = open('/home/B180093/anaconda3/envs/telechips/main/static/download/' + sdk_name[:-1] + '_' + data['patch_version'] + '_' + patch_description + '.' + f1_ext, 'rb')
    os.remove(os.path.join('/home/B180093/anaconda3/envs/telechips/main/static/download/', sdk_name[:-1] + '_' + data['patch_version'] + '_' + patch_description + '.' + f1_ext))
    f1.close()

    f2 = request.files['releaseNote']
    f2_ext = f2.filename.split('.')[1]
    f2.save('/home/B180093/anaconda3/envs/telechips/main/static/download/' + sdk_name[:-1] + '_' + data['patch_version'] + '_' + patch_description + '.' + f2_ext)
    files2 = open('/home/B180093/anaconda3/envs/telechips/main/static/download/' + sdk_name[:-1] + '_' + data['patch_version'] + '_' + patch_description + '.' + f2_ext, 'rb')
    os.remove(os.path.join('/home/B180093/anaconda3/envs/telechips/main/static/download/', sdk_name[:-1] + '_' + data['patch_version'] + '_' + patch_description + '.' + f2_ext))
    f2.close()

    url_cssapi = 'https://css.telechips.com/api/patch/save'

    headers = {
        'accept': 'application/json;charset=UTF-8',
        'x-token': '123qwe123'
    }

    if data['deploy_type'] == '1':
        tims_idx = 0
    else:
        tims_idx = int(data['customer_project'])

    reporter = data['reporter'].split('_')
    upload = {"patchCode": files1, "releaseNote": files2}

    if data['patch_type'] == 'TIMS (고객사 이슈)':
        patch_type = '1'
        send_data = {
                "pageIdx": data['sdkIdx'].split('_')[1],
                "sdkIdx": data['sdkIdx'].split('_')[0],
                "patchDeploymentType": data['deploy_type'],
                "patchDescription": data['patch_des'],
                "patchNumber": data['patch_version'],  # 중복되면 안됨
                "patchType": patch_type,
                'patchTimsIdx': data['project_id'],
                'patchTimsProjectKey': data['issue_key'].split('-')[0],
                'patchImIssueId': data['issue_key'].split('-')[-1],
                'imIssueId': data['customer_key'],
                'timsIdx': tims_idx,
                "regId": reporter[2],
                "regName": reporter[1],
                "regDept": reporter[0],
                "comment": data['comment'],
                'internal ': True
            }

    elif data['patch_type'] == 'TCS (QA 검증 이슈)':
        patch_type = '2'
        send_data = {
                "pageIdx": data['sdkIdx'].split('_')[1],
                "sdkIdx": data['sdkIdx'].split('_')[0],
                "patchDeploymentType": data['deploy_type'],
                "patchDescription": data['patch_des'],
                "patchNumber": data['patch_version'],  # 중복되면 안됨
                "patchType": patch_type,
                'patchTcsIdx': data['project_id'],
                'patchTcsProjectKey': data['issue_key'].split('-')[0],
                'patchQaIssueId': data['issue_key'].split('-')[-1],
                'imIssueId': data['customer_key'],
                'timsIdx': tims_idx,
                "regId": reporter[2],
                "regName": reporter[1],
                "regDept": reporter[0],
                "comment": data['comment'],
                'internal ': True
            }

    else:
        patch_type = '6'
        send_data = {
                "pageIdx": data['sdkIdx'].split('_')[1],
                "sdkIdx": data['sdkIdx'].split('_')[0],
                "patchDeploymentType": data['deploy_type'],
                "patchDescription": data['patch_des'],
                "patchNumber": data['patch_version'],  # 중복되면 안됨
                "patchType": patch_type,
                'imIssueId': data['customer_project'],
                'timsIdx': tims_idx,
                "regId": reporter[2],
                "regName": reporter[1],
                "regDept": reporter[0],
                "comment": data['comment'],
                'internal ': True
            }

    request_css = requests.post(url_cssapi, headers=headers, files=upload, data=send_data)
    result_css = json.loads(request_css.text)

    if 'data' in result_css.keys():
        result_html = '<html><body>' + result_css['msg'] + '<br><a href="' + 'https://css.telechips.com/user/patch/view/' + str(result_css['data']['idx']) + '?pageIdx=' + str(result_css['data']['pageIdx']) + '" target="_blank">Patch Link (CSS에 로그인이 되어 있어야 정상적인 이동이 가능합니다.)' + '</a></body></html>'

    elif 'msg' in result_css.keys():
        result_html = result_css['msg']

    else:
        result_html = 'Error 신호찬M에게 문의'

    files1.close()
    files2.close()

    return result_html


@blueprint.route('/issue', methods=['GET', 'POST'])
def issue():  # Git Commit Message 생성할때 Error 발생하면 이동되는 페이지
    return render_template("issue.html")


@blueprint.route('/download', methods=['GET', 'POST'])
def download():  # Git Commit Message 생성 후 Release Note 다운로드 하는 함수
    try:
        note_data = request.form

        document = Document('main/static/release_note.docx')
        style = document.styles['Normal']
        style.font.name = 'Tahoma'

        sdk_name = note_data['sdk_name']
        sdk_version = note_data['sdk_version']
        board = note_data['evb']
        patch_name = note_data['patch_name']
        patch_version = note_data['patch_version']
        symptom = note_data['symptom']
        cause = note_data['cause']
        solution = note_data['solution']
        pre = note_data['pre_patch']
        how = note_data['how_patch']

        symptom = note_data['symptom']

        p = document.paragraphs[1].add_run(sdk_name + '_' + sdk_version + '\n' + note_data['patch_name'])
        p.bold = True
        p.font.color.rgb = RGBColor(255, 255, 255)
        p.font.size = Pt(20)

        for i in document.paragraphs:
            if '{sdk_name}' in i.text:
                i.text = i.text.replace('{sdk_name}', sdk_name)

            elif '{sdk_version}' in i.text:
                i.text = i.text.replace('{sdk_version}', sdk_version)

            elif '{board}' in i.text:
                i.text = i.text.replace('{board}', board)

            elif '{patch_name}' in i.text:
                i.text = i.text.replace('{patch_name}', patch_name)

            elif '{patch_version}' in i.text:
                i.text = i.text.replace('{patch_version}', patch_version)

            elif '{symptom}' in i.text:
                i.text = i.text.replace('{symptom}', symptom.replace('\r', ''))

            elif '{cause}' in i.text:
                i.text = i.text.replace('{cause}', cause.replace('\r', ''))

            elif '{solution}' in i.text:
                i.text = i.text.replace('{solution}', solution.replace('\r', ''))

            elif '{pre}' in i.text:
                i.text = i.text.replace('{pre}', pre.replace('\r', ''))

            elif '{how}' in i.text:
                i.text = i.text.replace('{how}', how.replace('\r', ''))

        document.save('main/static/result/Release_note.docx')

        with open(os.path.join('main/static/result/', 'Release_note.docx'), 'rb') as f:
            data = f.readlines()
        os.remove(os.path.join('main/static/result/', 'Release_note.docx'))

        filename = sdk_name + '_' + sdk_version + '_' + patch_version + '_' + patch_name + '.docx'

        return Response(data, headers={
            'Content-Type': 'application/msword',
            'Content-Disposition': 'attachment; filename=%s;' % filename})

    except Exception:
        return render_template("issue.html")


@blueprint.route('/download_pdf', methods=['GET', 'POST'])
def download_pdf():  # Git Commit Message 생성 후 Release Note 다운로드 하는 함수
    try:
        note_data = request.form

        document = Document('main/static/release_note.docx')
        style = document.styles['Normal']
        style.font.name = 'Tahoma'

        sdk_name = note_data['sdk_name']
        sdk_version = note_data['sdk_version']
        board = note_data['evb']
        patch_name = note_data['patch_name']
        patch_version = note_data['patch_version']
        symptom = note_data['symptom']
        cause = note_data['cause']
        solution = note_data['solution']
        pre = note_data['pre_patch']
        how = note_data['how_patch']

        symptom = note_data['symptom']

        p = document.paragraphs[1].add_run(sdk_name + '_' + sdk_version + '\n' + note_data['patch_name'])
        p.bold = True
        p.font.color.rgb = RGBColor(255, 255, 255)
        p.font.size = Pt(20)

        for i in document.paragraphs:
            if '{sdk_name}' in i.text:
                i.text = i.text.replace('{sdk_name}', sdk_name)

            elif '{sdk_version}' in i.text:
                i.text = i.text.replace('{sdk_version}', sdk_version)

            elif '{board}' in i.text:
                i.text = i.text.replace('{board}', board)

            elif '{patch_name}' in i.text:
                i.text = i.text.replace('{patch_name}', patch_name)

            elif '{patch_version}' in i.text:
                i.text = i.text.replace('{patch_version}', patch_version)

            elif '{symptom}' in i.text:
                i.text = i.text.replace('{symptom}', symptom.replace('\r', ''))

            elif '{cause}' in i.text:
                i.text = i.text.replace('{cause}', cause.replace('\r', ''))

            elif '{solution}' in i.text:
                i.text = i.text.replace('{solution}', solution.replace('\r', ''))

            elif '{pre}' in i.text:
                i.text = i.text.replace('{pre}', pre.replace('\r', ''))

            elif '{how}' in i.text:
                i.text = i.text.replace('{how}', how.replace('\r', ''))

        document.save('main/static/result/Release_note.docx')
        subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'main/static/result/Release_note.docx', '--outdir', 'main/static/result'])

        with open(os.path.join('main/static/result', 'Release_note.pdf'), 'rb') as f:
            data = f.readlines()
        os.remove(os.path.join('main/static/result/', 'Release_note.docx'))
        os.remove(os.path.join('main/static/result', 'Release_note.pdf'))

        filename = sdk_name + '_' + sdk_version + '_' + patch_version + '_' + patch_name + '.pdf'

        return Response(data, headers={
            'Content-Type': 'application/msword',
            'Content-Disposition': 'attachment; filename=%s;' % filename})

    except Exception:
        return render_template("issue.html")


@blueprint.route('/git', methods=['GET', 'POST'])
def git():  # Git Commit Message 생성을 위한 최초 페이지
    conn = pymongo.MongoClient("219.240.43.93", 27017)
    db = conn.tcs

    col = db.chip_list
    data = list(col.find({}, {"_id": 0}))
    chip_list = sorted(data, reverse=False, key=lambda x: (x['sub_chip']))

    col = db.project_key_category
    data1 = list(col.find({"projectcategory": "QA 프로젝트"}, {"_id": 0, "key": -1, "name": -1}))
    project_list = sorted(data1, reverse=False, key=lambda x: (x['name']))

    col = db.user_data
    data2 = list(col.find({}, {"_id": 0, "name": -1, "employee_No": -1}))
    user_list = sorted(data2, reverse=False, key=lambda x: (x['name']))

    return render_template("git.html", chip=chip_list, project=project_list, user=user_list)


@blueprint.route('/result', methods=['GET', 'POST'])
def result():  # Git Commit Message 생성 페이지에서 버튼을 누르면 이동되는 페이지
    try:
        data = request.form

        day = datetime.now()
        date = day.strftime('%Y-%m-%d')

        conn = pymongo.MongoClient("219.240.43.93", 27017)
        db = conn.tcs
        col = db.id_pw
        pw_data = col.find({})
        id_pw = {'os_username': pw_data[1]['id'], 'os_password': pw_data[1]['pw']}

        project_name = ''
        chip = 'Common'
        os = 'Common'

        # QA프로젝트 정보 구하기
        r = requests.get('https://tcs.telechips.com:8443/rest/profields/api/2.0/values/projects/' + data['project_key'], auth=(pw_data[1]['id'], pw_data[1]['pw']))
        profield_data = json.loads(r.text)

        if isinstance(profield_data, list):
            for i in range(0, len(profield_data)):
                if profield_data[i]['field']['id'] == -2 and 'value' in profield_data[i].keys():
                    project_name = profield_data[i]['value']['value'].split(']')[1].replace(' ', '')
                elif profield_data[i]['field']['id'] == 43 and 'value' in profield_data[i].keys():
                    chip = profield_data[i]['value']['value']['value']
                elif profield_data[i]['field']['id'] == 49 and 'value' in profield_data[i].keys():
                    os = profield_data[i]['value']['value']['value']
        elif isinstance(profield_data, dict):
            project_name = ''
            chip = ''
            os = ''

        headers = {'Content-Type': 'application/json'}

        summary = ''
        patch_type_issuekey = ''
        if data['patch_type'] == 'tims':
            summary = '[Common][' + data['tims_issue'] + ']' + data['commit_message']
            patch_type = 'TIMS (고객사 이슈)'
            patch_type_issuekey = data['tims_issue']
        if data['patch_type'] == 'qa':
            summary = '[Common][' + data['qa_issue'] + ']' + data['commit_message']
            patch_type = 'TCS (QA 검증 이슈)'
            patch_type_issuekey = data['qa_issue']
        if data['patch_type'] == 'rnd':
            summary = '[Common][RND]' + data['commit_message']
            patch_type = 'RND (신규 기능 추가/개선)'
            patch_type_issuekey = 'None'

        error_result = ''
        key_result = ''

        issue_data1 = {
                    'fields': {
                                'project': {'key': data['project_key']},
                                'summary': summary,
                                'reporter': {'name': data['reporter']},
                                'assignee': {'name': data['assignee']},
                                'issuetype': {'name': 'Patch'},
                                'customfield_10200': date,  # start date
                                'duedate': date,
                                'customfield_11101': {'value': chip},  # chip
                                'customfield_11726': {'value': patch_type},  # Patch Type
                                'customfield_12001': patch_type_issuekey,  # Patch Type issuekey
                                'customfield_10684': {'value': os},  # O/S
                                'customfield_10686': data['sdk_version'],  # SDK Version
                                'customfield_10685': data['patch_version'],  # Patch Version
                                'customfield_10689': {'value': data['sub_device']},  # Sub Device/s
                                'customfield_11704': data['cause'],  # 이슈발생원인
                                'customfield_11705': data['symptom'],  # 이슈재현방법
                                'customfield_11709': {'value': data['evb']},  # Test EVB
                                'customfield_11706': data['solution'],  # 이슈해결방법
                                'customfield_11707': data['test_result'],  # 이슈처리결과
                                'customfield_11708': 'Push 후 내용 입력 예정입니다.'  # Branch/ Commit info
                            }
                }

        # TCS Issue 생성
        r = requests.post('https://tcs.telechips.com:8443/rest/api/2/issue', data=json.dumps(issue_data1), headers=headers, auth=(id_pw['os_username'], id_pw['os_password']))
        if 'errors' in r.json().keys():
            error_result = str(r.json()['errors'])
        if 'key' in r.json().keys():
            key_result = r.json()['key']

        data1 = {'project_name': project_name, 'chip': chip, 'os': os, 'date': date, 'error_result': error_result, 'key_result': key_result}

        return render_template("result.html", data=data, data1=data1)

    except Exception as ex:
        print("에러가 발생했습니다. RIT 신호찬M에게 문의 부탁드립니다.", ex)
